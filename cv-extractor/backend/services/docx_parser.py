from docx import Document
from typing import Optional


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extrait le texte brut d'un fichier DOCX (Word).
    
    Args:
        file_path: Chemin vers le fichier DOCX
        
    Returns:
        Le texte extrait ou None si erreur
        
    Raises:
        Exception: Si le fichier ne peut pas être lu
    """
    try:
        # Ouvre le document Word
        doc = Document(file_path)
        
        text = ""
        
        # Parcourt tous les paragraphes du document
        for paragraph in doc.paragraphs:
            # Ajoute le texte de chaque paragraphe
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extrait aussi le texte des tableaux (si présents)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        # Vérifie qu'on a bien extrait du texte
        if not text.strip():
            raise ValueError("Le fichier DOCX ne contient pas de texte")
            
        return text
        
    except Exception as e:
        raise Exception(f"Erreur lors de l'extraction du DOCX : {str(e)}")


def clean_text(text: str) -> str:
    """
    Nettoie le texte extrait :
    - Supprime les espaces multiples
    - Supprime les sauts de ligne excessifs
    - Normalise les espaces
    
    Args:
        text: Le texte brut à nettoyer
        
    Returns:
        Le texte nettoyé
    """
    if not text:
        return ""
    
    # Remplace les multiples espaces par un seul
    text = " ".join(text.split())
    
    # Remplace les multiples sauts de ligne par un seul
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    text = "\n".join(lines)
    
    return text